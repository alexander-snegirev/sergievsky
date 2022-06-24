import json
import sys

from docx import Document


def _handler_get_table_data(tables):
    schedule_list = []
    for row in tables.rows[1:]:
        date, to_whom, time, type_worship = row.cells
        if date.text != '':
            if len(schedule_list) > 0 and date.text == schedule_list[-1]['date']:
                schedule_list[-1]['tt'].append(dict(time=time.text, type_worship=type_worship.text))
            else:
                schedule = dict(
                    date=date.text,
                    to_whom=to_whom.text,
                    tt=[dict(time=time.text, type_worship=type_worship.text)]
                )
                schedule_list.append(schedule)
    return schedule_list


def get_data(docx_file):
    document = Document(docx_file)
    paragraphs = ([p.text for p in document.paragraphs if p.text != ''])
    data = {}
    for item in range(len(paragraphs)):
        parsing_data = dict(data=_handler_get_table_data(document.tables[item]))
        # data_raw = {item: {paragraphs[item]: {}}}
        data_raw = {item: {'title_table': paragraphs[item], 'tables': {}}}
        data_raw[item]['tables'].update(parsing_data)
        data.update(data_raw)
    return json.dumps(data, indent=2, ensure_ascii=False)


if __name__ == '__main__':
    file = sys.argv[-1]
    get_data(file)
